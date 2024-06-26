#Importing dependencies
import pyautogui
import time
from PIL import Image, ImageEnhance
import pytesseract
import ctypes
import re
import openpyxl
import shutil
import csv
import os
from datetime import date
from datetime import timedelta
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


# Round Order toggle. Set to false to revert to alphabetical room order for list sorting
# When set to true, list order starts with the higher floors and works its way down, with list order for 3B/4B reversed
ROUND_ORDER = True

# Vista font size selection
FONT_SIZE = "12"

# Get current logged in windows user. This is used as a workaround for user-only installs of tesseract
USER = os.getlogin()

# This command hides the window for the program when after it starts to not cover the Vista window
ctypes.windll.user32.ShowWindow( ctypes.windll.kernel32.GetConsoleWindow(), 6 )

# Whether yesterday's excel file is present or not
YESTERDAY_PRESENT = False

# Workaround for common OCR errors when reading patient's rooms
def roomClear(x):
    print(x)
    x=x.replace("2-","2A-")
    x=x.replace("28-","2B-")
    x=x.replace("A0","A-0")
    x=x.replace("B0","B-0")
    x=x.replace("C1","C-1")
    x=x.replace("AT","A-1")
    if x[0]=="S": x="3"+x[1:]
    if x[0]=="A": x="4"+x[1:]
    if x[0]=="B": x="5"+x[1:]
    if x[1]=="4": x=x[0]+"A"+x[2:]
    if x[0:2]=="33": x=x.replace("33","3")
    if x[0:2]=="44": x=x.replace("44","3")
    if x[0:2]=="55": x=x.replace("55","3")
    x=x.replace("4AC","4C")
    x=x.replace("5BC","5C")
    if x[2]=="0": x=x[0:2]+"-"+x[3:]
    if x[2]==" ": x=x[0:2]+"-"+x[3:]
    x=x.replace("A1","A-1")
    x=x.replace("B1","B-1")
    if x[-1]=="4": x=x[:-1]+"A"
    if x[-1]=="8": x=x[:-1]+"B"
    x=x.replace("3H","9H")
    x=x.replace("AA","A")
    x=x.replace("-A","-4")
    x=x.replace("-O","-0")
    x=x.replace("-2A","-24")
    x=x.replace("-2B","-28")
    x=x.strip("-")
    x=x.replace("5LD-TX","BLD-TX")
    x=x.replace("DAYCAS","DAYCASE")
    x=x.replace("DC-KSBi","DC-KSB")
    x=x.replace("3-","3A-")
    x=x.replace("4-","4A-")
    x=x.replace("4A4A","4A")
    x=x.replace("40-","4C-")
    x=x.replace("5-","5C-")
    x=x.replace("QH","9H")
    x=x.replace("A2","A-2")
    x=x.replace("0-","C-")
    print(x)
    return x

# Workaround for common OCR errors when reading patient's MRN
# Known error: Sometimes the number '5' is read as '6'. This issue is unavoidable with current implementation.
#              Possible fixes require image manipulation which will increase time for list generation significantly
def mrnClear(x):
    print(x)
    x=x.strip().strip("\'\"|/\\").strip("\'\"|/\\‘°").replace(" ","")
    x=x.replace("i","1")
    x=x.replace("o","9")
    x=x.replace("D","5")
    x=x.replace("?","7")
    if not x or x.isspace(): return ""
    if x[0]=="0": x="5"+x
    x=x.strip("!")
    if x[0]=="T": x="1"+x[1:]
    if len(x)==7: x="2"+x[2:]
    print(x)
    return x

# Given a patient, this function returns the floor the patient is on
def getFloor(x):
    dash=x.room.find("-")
    if dash==-1: dash=2
    return x.room[0:dash]
    
# Given a patient, this function returns the room number the patient is in (without the floor)
def getRoom(x):
    dash=x.room.find("-")
    if dash==-1: dash=2
    return x.room[dash:].strip("-")

# Adjusts offsets for screenshots based on font size
def normalize(x):
    return int(x*int(FONT_SIZE)/12)

# Patient class definition. This class mostly functions as a data container and to facilitate list sorting.
class Patient:
    # Class constructor
    def __init__(self, name, room, mrn):
        self.name = name.replace(",,",",")
        self.room = room
        self.mrn = mrn
        self.new = False
    
    # The "less than" comparator definition for this class. Python requires defining at least one comparision method for sorting objects.
    # This function works by defining what the "<" operator returns  (i.e. when patient a (self)< patient b (obj) is true).
    # Two sorting methods are presented based on the ROUND_ORDER value as defined above. Both methods require a special case for floors 10 and 11 
    # as basic string comparison works one character at a time and both will be read as 1* (i.e. less than 2)
    def __lt__(self,obj):
        # Rounding order sorting method
        if (ROUND_ORDER):
            if(self.room[0:2]=="11" and not obj.room[0:2]=="11"): return True
            if(not self.room[0:2]=="11" and obj.room[0:2]=="11"): return False
            if(self.room[0:2]=="10" and not obj.room[0:2]=="10"): return True
            if(not self.room[0:2]=="10" and obj.room[0:2]=="10"): return False
            if(getFloor(self)[0]==getFloor(obj)[0]):
                if(getFloor(self)[1]=="C" and not getFloor(obj)[1]=="C"): return True
                if(not getFloor(self)[1]=="C" and getFloor(obj)[1]=="C"): return False
                if(getFloor(self)[1]=="B" and getFloor(obj)[1]=="A"): return False
                if(getFloor(self)[1]=="A" and getFloor(obj)[1]=="B"): return True
            if(getFloor(self)<getFloor(obj)): return False
            if(getFloor(self)>getFloor(obj)): return True
            if(getFloor(self)==getFloor(obj)):
                if(getFloor(self)[0:2]=="4B"): return getRoom(self)>getRoom(obj)
                if(getFloor(self)[0:2]=="3B"): return getRoom(self)>getRoom(obj)
                else:return getRoom(self)<getRoom(obj)
        # Traditional sorting method
        else:
            if(self.room[0:2]=="11" and not obj.room[0:2]=="11"): return False
            if(not self.room[0:2]=="11" and obj.room[0:2]=="11"): return True
            if(self.room[0:2]=="10" and not obj.room[0:2]=="10"): return False
            if(not self.room[0:2]=="10" and obj.room[0:2]=="10"): return True
            return ((self.room) < (obj.room))
 
# Initialization for pyautogui (the input automation library) and pytesseract (the OCR library) 
pyautogui.FAILSAFE = True
pytesseract.pytesseract.tesseract_cmd=fr'C:\Users\{USER}\AppData\Local\Programs\Tesseract-OCR\tesseract.exe'

# Using screenshots to find the "Cancel" button
cancelTop=pyautogui.locateOnScreen(fr"Data\{FONT_SIZE}\cancel.png").top
cancelLeft=pyautogui.locateOnScreen(fr"Data\{FONT_SIZE}\cancel.png").left

# Defining a dictionary containing the relavent team
team=[]

scrollbox=pyautogui.locateOnScreen(fr"Data\{FONT_SIZE}\patientScroll.png") # The scroll bar for patient list
# Specifies an area to look for data relative to the cancel button (for optimization)
cancel=pyautogui.locateOnScreen(fr"Data\{FONT_SIZE}\cancel.png")
searchregion=(cancel.left-600,cancel.top,700,400)

# Clicks on the first patient in the list. Location is relative to the scroll bar 
pyautogui.click(scrollbox.left-250,scrollbox.top+5)
time.sleep(0.8)
# Finds region for screenshot of patient data in vista
ssn=pyautogui.locateOnScreen(fr"Data\{FONT_SIZE}\ssn.png", region=searchregion)
nameRegion=(ssn.left,ssn.top-normalize(20),normalize(400),normalize(20))
khcc=pyautogui.locateOnScreen(fr"Data\{FONT_SIZE}\khccroom.png", region=searchregion)
roomRegion=(khcc.left,khcc.top+normalize(20),normalize(75),normalize(20))
mrn=pyautogui.locateOnScreen(fr"Data\{FONT_SIZE}\mrn.png", region=searchregion)
mrnRegion=(mrn.left+normalize(45),mrn.top,normalize(65),normalize(20))
# This assumes a maximum of 45 patients per team
for i in range(45):
    time.sleep(0.8)
    # Takes screenshots of name, MRN, and room and uses tesseract for OCR
    nameImage=pyautogui.screenshot(r'1.png', region=nameRegion)
    roomImage=ImageEnhance.Contrast(pyautogui.screenshot(r'3.png', region=roomRegion)).enhance(20)
    mrnImage=ImageEnhance.Contrast(pyautogui.screenshot(r'2.png', region=mrnRegion)).enhance(20)
    name=pytesseract.image_to_string(nameImage).strip().replace(".",",").replace("_",",").replace("|","").strip("\'\"|/\\").strip("\'\"|/\\")
    if not name or name.isspace(): 
        break
    mrn=mrnClear(pytesseract.image_to_string(mrnImage, config="--psm 7"))        
    roomstr=pytesseract.image_to_string(roomImage, lang="eng", config="--psm 7").strip().strip("\'\"|/\\,.").strip("\'\"|/\\").strip()
    # Code for skipping fake patients
    if not mrn or mrn.isspace() or "Test" in name or "Pacs" in name or not roomstr or len(roomstr)<=2:
        pyautogui.press('down')
        continue
    room=roomClear(roomstr)
    # Creating a new patient with the collected information, 
    # then pressing the down button to get next patient
    team.append(Patient(name,room,mrn))
    pyautogui.press('down')

team.sort()
# Creates new word document and sets some style formatting
document = Document()
style = document.styles['No Spacing']
font = style.font
font.name = 'Calibri'
font.size = Pt(16)
font.bold = True
normalstyle = document.styles['Normal']
normalstyle.font.name = 'Calibri'

title=document.add_paragraph(f'{team[0]} ({date.today()})')
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.style = style
row_number=len(team)+1
row_height=int(Inches(8.4)/row_number)
table=document.add_table(rows=row_number, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'
table.rows[0].cells[3].text="MRN"
table.rows[0].cells[2].text="Room"
table.rows[0].cells[1].text="Name"
i=1
for patient in team:
    table.rows[i].height=row_height
    table.rows[i].cells[3].text=patient.mrn
    table.rows[i].cells[3].width=0
    table.rows[i].cells[2].text=patient.room
    table.rows[i].cells[2].width=0
    table.rows[i].cells[1].text=patient.name 
    table.rows[i].cells[1].width=Cm(50)
    table.rows[i].cells[0].text=str(i) +("*" if patient.new else "")
    table.rows[i].cells[0].width=0
    # Adds shading if its a new patient
    if patient.new: 
        for cell in table.rows[i].cells:
            shading_elm = parse_xml(r'<w:shd {} w:fill="E9E9E9"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading_elm)
    i+=1
    
# Saves the word file. If an error occurs (Usually because word file is already open)
# it prompts the user to close it
while True:
    try:
        document.save(f"{date.today()}-Team.docx")
        break
    except:
        pyautogui.alert(text='Please close Word file and press OK', title='File write error', button='OK')    

