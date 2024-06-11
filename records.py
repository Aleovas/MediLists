#Importing dependencies
import pyautogui
import time
from PIL import Image, ImageEnhance
import pytesseract
import ctypes
import re
import openpyxl
import shutil
import csv
import os
from datetime import date
from datetime import timedelta
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import tkinter as tk
import tkinter.ttk as ttk
import webbrowser

# Round Order toggle. Set to false to revert to alphabetical room order for list sorting
# When set to true, list order starts with the higher floors and works its way down, with list order for 3B/4B reversed
ROUND_ORDER = True

# Vista font size selection
FONT_SIZE = "12"

# Get current logged in windows user. This is used as a workaround for user-only installs of tesseract
USER = os.getlogin()

# This command hides the window for the program when after it starts to not cover the Vista window
ctypes.windll.user32.ShowWindow( ctypes.windll.kernel32.GetConsoleWindow(), 6 )

# Whether yesterday's excel file is present or not
YESTERDAY_PRESENT = False

# Workaround for common OCR errors when reading patient's rooms
def roomClear(x):
    print(x)
    x=x.replace("2-","2A-")
    x=x.replace("28-","2B-")
    x=x.replace("A0","A-0")
    x=x.replace("B0","B-0")
    x=x.replace("C1","C-1")
    x=x.replace("AT","A-1")
    if x[0]=="S": x="3"+x[1:]
    if x[0]=="A": x="4"+x[1:]
    if x[0]=="B": x="5"+x[1:]
    if x[1]=="4": x=x[0]+"A"+x[2:]
    if x[0:2]=="33": x=x.replace("33","3")
    if x[0:2]=="44": x=x.replace("44","3")
    if x[0:2]=="55": x=x.replace("55","3")
    x=x.replace("4AC","4C")
    x=x.replace("5BC","5C")
    if x[2]=="0": x=x[0:2]+"-"+x[3:]
    if x[2]==" ": x=x[0:2]+"-"+x[3:]
    x=x.replace("A1","A-1")
    x=x.replace("B1","B-1")
    if x[-1]=="4": x=x[:-1]+"A"
    if x[-1]=="8": x=x[:-1]+"B"
    x=x.replace("3H","9H")
    x=x.replace("AA","A")
    x=x.replace("5H","8H")
    x=x.replace("-A","-4")
    x=x.replace("-O","-0")
    x=x.replace("-2A","-24")
    x=x.replace("-2B","-28")
    x=x.strip("-")
    x=x.replace("5LD-TX","BLD-TX")
    x=x.replace("DAYCAS","DAYCASE")
    x=x.replace("DC-KSBi","DC-KSB")
    x=x.replace("3-","3A-")
    x=x.replace("4-","4A-")
    x=x.replace("4A4A","4A")
    x=x.replace("40-","4C-")
    x=x.replace("5-","5C-")
    x=x.replace("QH","9H")
    x=x.replace("A2","A-2")
    x=x.replace("0-","C-")
    x=x.replace("T1","11")
    x=x.replace("1O","10")
    x=x.replace("TO","10")
    x=x.replace("1H-0","1H-")
    x=x.replace("0H-0","0H-")
    x=x.replace("(","1")
    x=x.replace(")","1")
    print(x)
    return x

# Workaround for common OCR errors when reading patient's MRN
# Known error: Sometimes the number '5' is read as '6'. This issue is unavoidable with current implementation.
#              Possible fixes require image manipulation which will increase time for list generation significantly
def mrnClear(x):
    print(x)
    x=x.strip().strip("\'\"|/\\").strip("\'\"|/\\‘°").replace(" ","")
    x=x.replace("i","1")
    x=x.replace("o","9")
    x=x.replace("D","5")
    x=x.replace("?","7")
    x=x.replace("S","8")
    if not x or x.isspace(): return ""
    if x[0]=="0": x="5"+x
    x=x.strip("!")
    if x[0]=="T": x="1"+x[1:]
    if len(x)==7: x="2"+x[2:]
    print(x)
    return x

# Given a patient, this function returns the floor the patient is on
def getFloor(x):
    if(x.room[0:2]=="ER"):return "00"
    dash=x.room.find("-")
    if dash==-1: dash=2
    return x.room[0:dash]
    
# Given a patient, this function returns the room number the patient is in (without the floor)
def getRoom(x):
    dash=x.room.find("-")
    if dash==-1: dash=2
    return x.room[dash:].strip("-")

# Adjusts offsets for screenshots based on font size
def normalize(x):
    return int(x*int(FONT_SIZE)/12)

# Patient class definition. This class mostly functions as a data container and to facilitate list sorting.
class Patient:
    # Class constructor
    def __init__(self, name, room, mrn):
        self.name = name.replace(",,",",")
        self.room = room
        self.mrn = mrn
        self.new = False
    
    # The "less than" comparator definition for this class. Python requires defining at least one comparision method for sorting objects.
    # This function works by defining what the "<" operator returns  (i.e. when patient a (self)< patient b (obj) is true).
    # Two sorting methods are presented based on the ROUND_ORDER value as defined above. Both methods require a special case for floors 10 and 11 
    # as basic string comparison works one character at a time and both will be read as 1* (i.e. less than 2)
    def __lt__(self,obj):
        # Rounding order sorting method
        if (ROUND_ORDER):
            if(not self.room.isnumeric() and obj.room.isnumeric()): return False
            if(self.room[0:2]=="11" and not obj.room[0:2]=="11"): return True
            if(not self.room[0:2]=="11" and obj.room[0:2]=="11"): return False
            if(self.room[0:2]=="10" and not obj.room[0:2]=="10"): return True
            if(not self.room[0:2]=="10" and obj.room[0:2]=="10"): return False
            if(getFloor(self)[0]==getFloor(obj)[0]):
                if(getFloor(self)[1]=="C" and not getFloor(obj)[1]=="C"): return True
                if(getFloor(self)[1]=="C" and not getFloor(obj)[1]=="C"): return True
                if(not getFloor(self)[1]=="C" and getFloor(obj)[1]=="C"): return False
                if(getFloor(self)[1]=="B" and getFloor(obj)[1]=="A"): return False
                if(getFloor(self)[1]=="A" and getFloor(obj)[1]=="B"): return True
            if(getFloor(self)<getFloor(obj)): return False
            if(getFloor(self)>getFloor(obj)): return True
            if(getFloor(self)==getFloor(obj)):
                if(getFloor(self)[0:2]=="4B"): return getRoom(self)>getRoom(obj)
                if(getFloor(self)[0:2]=="3B"): return getRoom(self)>getRoom(obj)
                else:return getRoom(self)<getRoom(obj)
        # Traditional sorting method
        else:
            if(self.room[0:2]=="11" and not obj.room[0:2]=="11"): return False
            if(not self.room[0:2]=="11" and obj.room[0:2]=="11"): return True
            if(self.room[0:2]=="10" and not obj.room[0:2]=="10"): return False
            if(not self.room[0:2]=="10" and obj.room[0:2]=="10"): return True
            return ((self.room) < (obj.room))

# Team class definition. This class also mostly functions as a data container 
class Team:
    # Class constructor
    def __init__(self, name, image):
        self.patients=[]
        self.name = name
        self.image = fr"Data\{FONT_SIZE}\{image}"
        self.counts = {"2A":0,"3A":0,"3B":0,"3C":0,"4A":0,"4B":0,"4C":0,"5C":0,"Towers":0}
        self.new = 0
        self.old = 0
        self.oldPatients=[]

# Initialization for pyautogui (the input automation library) and pytesseract (the OCR library) 
pyautogui.FAILSAFE = True
while true:
    if os.path.isfile(fr'C:\Users\{USER}\AppData\Local\Programs\Tesseract-OCR\tesseract.exe'):
        pytesseract.pytesseract.tesseract_cmd=fr'C:\Users\{USER}\AppData\Local\Programs\Tesseract-OCR\tesseract.exe'
        break
    else:
        pyautogui.alert(text='The OCR program Tesseract was not found. Please press OK to download it and then install it', title='Tesseract not found', button='OK')
        webbrowser.open("https://digi.bib.uni-mannheim.de/tesseract/tesseract-ocr-w64-setup-5.4.0.20240606.exe")
        pyautogui.alert(text='Please click OK after installing tesseract', title='Tesseract not found', button='OK')


# Defining the teams that will be searched for patients
teams={
    "Team 1":Team("Team 1", "team1.png"),
    "Team 2":Team("Team 2", "team2.png"), 
    "Team 3":Team("Team 3", "team3.png"),
    "Team 4":Team("Team 4", "team4.png"),
    "Team 5":Team("Team 5", "team5.png"),
    "Leukemia":Team("Leukemia", "teamLeukemia.png"),
    "Lymphoma":Team("Lymphoma", "teamLymphoma.png"),
    "Palliative":Team("Palliative", "teampal1.png") 
}

skip=["Team 4"]

yesterdayTeams={
    "Team 1":[],
    "Team 2":[],
    "Team 3":[],
    "Team 5":[],
    "Lymphoma":[],
    "Leukemia":[],
    "Palliative":[]
}

personal={}


def transferPatients():
    # Getting the override list 
    override={}
    with open('override.csv') as overrideCSV:
        csv_reader = csv.DictReader(overrideCSV)
        for row in csv_reader:
            override[row["mrn"]]=row["team"]

    # Moving patients in override list, and removing patients in certain floors
    ignoredFloors=['2B','4H','5H','6H']        
    for team in teams.values():
        remove=[]
        for patient in team.patients:
            if patient.mrn in override.keys() and not override[patient.mrn]==team[0]:
                if not override[patient.mrn]=="NA": teams[override[patient.mrn]].patients.append(patient)
                remove.append(patient)
            if patient.room[0:2] in ignoredFloors:
                remove.append(patient)
            if "PICU" in patient.room: remove.append(patient)
        for patient in remove:
            try:team.patients.remove(patient)
            except: 
                print(str(patient))

def writeExcel():
    # Copies the template excel file to new file named after today's date
    # Then it gets the team counts sheet and finds the relavent teams
    excelName= f'{date.today()}.xlsx'
    while True:
        try:
            shutil.copy("template.xlsx",excelName)
            break
        except:
            pyautogui.alert(text='Please close Excel file and press OK', title='File write error', button='OK')
    wb = openpyxl.load_workbook(excelName)
    counts = wb['Counts']
    teamNameCells=tuple(counts["A2":"A8"])
    # Updates team count
    for team in teams.values():
        if team.name in skip:continue
        for cell in teamNameCells:
            if cell[0].value == team.name:
                r=cell[0].row
                c=2
                for count in team.counts.values():
                    counts.cell(row=r,column=c).value=count
                    c+=1
                counts.cell(row=r,column=12).value=team.old
                counts.cell(row=r,column=13).value=team.new

        # Gets sheet for team list. Adds patients to each team's list
        teamSheet=[team.name]
        r=2
        for patient in team.patients:
            teamSheet.cell(row=r,column=1).value=str(r-1)+("+"if patient.new else "")
            teamSheet.cell(row=r,column=2).value=patient.name
            teamSheet.cell(row=r,column=3).value=patient.room
            teamSheet.cell(row=r,column=4).value=patient.mrn
            r+=1
        r=2
        teamSheet.cell(row=1,column=8).value="Discharge/transfer"
        teamSheet.cell(row=1,column=8).value="No."
        teamSheet.cell(row=1,column=9).value="Name"
        teamSheet.cell(row=1,column=10).value="Room"
        teamSheet.cell(row=1,column=11).value="MRN"
        for patient in team.oldPatients:
            teamSheet.cell(row=r,column=8).value=str(r-1)
            teamSheet.cell(row=r,column=9).value=patient.name
            teamSheet.cell(row=r,column=10).value=patient.room
            teamSheet.cell(row=r,column=11).value=patient.mrn

    # Saves the excel file. If an error occurs (Usually because excel file is already open)
    # it prompts the user to close it
    for i in range(10):
        try:
            wb.save(excelName)
            wb.close()
            break
        except:
            pyautogui.alert(text='Please close Excel file and press OK', title='File write error', button='OK')
        

def updateCounts():
    for team in teams.values():
        for key in team.counts.keys():
            team.counts[key]=0
        team.patients.sort()
        for patient in team.patients:
            # If patient's room isn't in floors 2-5, adds them to the towers count 
            # which acts as a catch-all for other patients (Daycase, IV unit, etc.)
            pattern=re.compile(r'[2-5][ABC]')
            if(pattern.match(patient.room[0:2])): team.counts[patient.room[0:2]]+=1
            else: team.counts["Towers"]+=1
        print(f'{team.name}: {team.counts}')

def comparePatients():
    # If yesterday's excel file is present it counts missing and new patients for each team
    if YESTERDAY_PRESENT:
        for team in teams.values():
            if team.name()=="Team 4": continue
            mrnsToday=[x.mrn for x in team.patients]
            mrnsYesterday=[x.mrn for x in yesterdayTeams[team.name]]
            for patient in team.patients:
                if patient.mrn not in mrnsYesterday:
                    patient.new=True
                    team.new += 1 # Updates new patient counter for the team
            for patient in yesterdayTeams[team.name]: # Updates missing patient counter
                if patient.mrn not in mrnsToday: 
                    team.old += 1
                    team.oldPatients.append(patient)
            print(f'New/Old: {team.new}/{team.old}')
                        

def writeWord():
    sheets={}
    for team in teams.values():
        sheets[team.name]=team.patients
    for team in personal.items():
        sheets[team[0]]=team[1][1]
    
    # Creates new word document and sets some style formatting
    document = Document()
    style = document.styles['No Spacing']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(16)
    font.bold = True
    normalstyle = document.styles['Normal']
    normalstyle.font.name = 'Calibri'

    # Creates list for each team in the word document
    for team in sheets:
        # When it gets to team 4, it ignores it and goes to other teams
        # Team 4 doesn't need to have a list
        if team[0] in skip: continue
        if len(team[1])==0: continue

        title=document.add_paragraph(f'{team[0]} ({date.today()})')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.style = style
        row_number=len(team[1])+1
        row_height=int(Inches(8.4)/row_number)
        table=document.add_table(rows=row_number, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        table.rows[0].cells[3].text="MRN"
        table.rows[0].cells[2].text="Room"
        table.rows[0].cells[1].text="Name"
        i=1
        for patient in team[1]:
            table.rows[i].height=row_height
            table.rows[i].cells[3].text=patient.mrn
            table.rows[i].cells[3].width=0
            table.rows[i].cells[2].text=patient.room
            table.rows[i].cells[2].width=0
            table.rows[i].cells[1].text=patient.name 
            table.rows[i].cells[1].width=Cm(50)
            table.rows[i].cells[0].text=str(i) +("+" if patient.new else "")
            table.rows[i].cells[0].width=0
            # Adds shading if its a new patient
            if patient.new: 
                for cell in table.rows[i].cells:
                    shading_elm = parse_xml(r'<w:shd {} w:fill="E9E9E9"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shading_elm)
            i+=1
        if team[0]!=list(sheets.keys())[-1]: document.add_page_break()   

    # Saves the word file. If an error occurs (Usually because word file is already open)
    # it prompts the user to close it
    while True:
        try:
            document.save(f"{date.today()}.docx")
            break
        except:
            pyautogui.alert(text='Please close Word file and press OK', title='File write error', button='OK') 


cancel=(0,0,10,10)
def getPatientList(palliative=False):
    scrollbox=pyautogui.locateOnScreen(fr"Data\{FONT_SIZE}\patientScroll.png") # The scroll bar for patient list
    # Specifies an area to look for data relative to the cancel button (for optimization)
    #cancel=pyautogui.locateOnScreen(fr"Data\{FONT_SIZE}\cancel.png")
    searchregion=(cancel.left-600,cancel.top,700,400)

    # Clicks on the first patient in the list. Location is relative to the scroll bar 
    pyautogui.click(scrollbox.left-250,scrollbox.top+5)
    pal=False
    time.sleep(0.8)
    # Finds region for screenshot of patient data in vista
    attempt_count=0
    nameRegion=(0,0,10,10)
    roomRegion=(0,0,10,10)
    mrnRegion=(0,0,10,10)
    while true:
        try:
            ssn=pyautogui.locateOnScreen(fr"Data\{FONT_SIZE}\ssn.png", region=searchregion)
            nameRegion=(ssn.left,ssn.top-normalize(20),normalize(400),normalize(20))
            khcc=pyautogui.locateOnScreen(fr"Data\{FONT_SIZE}\khccroom.png", region=searchregion)
            roomRegion=(khcc.left,khcc.top+normalize(20),normalize(75),normalize(20))
            mrn=pyautogui.locateOnScreen(fr"Data\{FONT_SIZE}\mrn.png", region=searchregion)
            mrnRegion=(mrn.left+normalize(45),mrn.top,normalize(65),normalize(20))
            attempt_count+=1
            break
        except:
            pyautogui.press('down')
            if attempt_count == 5: break
            continue
    patients=[]
    # This assumes a maximum of 45 patients per team
    for i in range(45):
        if attempt_count >= 3: break
        time.sleep(0.8)
        # Takes screenshots of name, MRN, and room and uses tesseract for OCR
        nameImage=pyautogui.screenshot(r'1.png', region=nameRegion)
        roomImage=ImageEnhance.Contrast(pyautogui.screenshot(r'3.png', region=roomRegion)).enhance(20)
        mrnImage=ImageEnhance.Contrast(pyautogui.screenshot(r'2.png', region=mrnRegion)).enhance(20)
        name=pytesseract.image_to_string(nameImage).strip().replace(".",",").replace("_",",").replace("|","").strip("\'\"|/\\").strip("\'\"|/\\")
        if not name or name.isspace(): 
            # Moving to next list if it reaches an empty name (i.e. end of list)
            # Below is special code to account for palliative being split into two lists 
            if pal: break
            if(palliative):
                box=pyautogui.locateOnScreen(fr"Data\{FONT_SIZE}\teampal2.png")
                pyautogui.click(box.left,box.top)
                time.sleep(1)
                pyautogui.click(scrollbox.left-300,scrollbox.top+5)
                pal=True
                continue
            break
        mrn=mrnClear(pytesseract.image_to_string(mrnImage, config="--psm 7"))        
        roomstr=pytesseract.image_to_string(roomImage, lang="eng", config="--psm 7").strip().strip("\'\"|/\\,.").strip("\'\"|/\\").strip()
        # Code for skipping fake patients
        if not mrn or mrn.isspace() or "Test" in name or "Pacs" in name or not roomstr or len(roomstr)<=2:
            pyautogui.press('down')
            attempt_count+=1
            continue
        attempt_count-=1
        room=roomClear(roomstr)
        # Creating a new patient with the collected information, 
        # then pressing the down button to get next patient
        patients.append(Patient(name,room,mrn))
        pyautogui.press('down')
    return patients
    
def getLastExcel():
    # If excel file is present for yesterday (or in the last 10 days), it reads it for extra statistics
    for i in range(1,11):
        yesterdayExcel=f'{date.today()-timedelta(days = i)}.xlsx'
        if os.path.isfile(yesterdayExcel):
            wb = openpyxl.load_workbook(yesterdayExcel)
            global YESTERDAY_PRESENT
            YESTERDAY_PRESENT = True
            for team in yesterdayTeams.items():
                teamSheet=wb.get_sheet_by_name(team[0])
                r=2
                while True:
                    name=teamSheet.cell(row=r,column=2).value
                    room=teamSheet.cell(row=r,column=3).value
                    mrn=teamSheet.cell(row=r,column=4).value
                    if not mrn or str(mrn).isspace(): break
                    team[1].append(Patient(name,room,mrn))           
                    r+=1
            wb.close()
            break

def getPersonal():
    with open('personal.csv') as personalCSV:
        csv_reader = csv.DictReader(personalCSV)
        for row in csv_reader:
            print(row)
            personal[row["name"]]=(row["vista_title"],[])
    pyautogui.click(pyautogui.locateCenterOnScreen(fr"Data\{FONT_SIZE}\personal.png"))
    time.sleep(0.8)
    for team in personal.items():
        team[1][1].clear()
        box=pyautogui.locateOnScreen(fr"Data\{FONT_SIZE}\teamSearchbar.png")
        pyautogui.click(box.left,box.top)
        pyautogui.typewrite(team[1][0])
        time.sleep(0.5)
        for patient in getPatientList():
            team[1][1].append(patient)
        team[1][1].sort()


def records():
    for i in range(10):
        try:
            global cancel
            cancel=pyautogui.locateOnScreen(fr"Data\{FONT_SIZE}\cancel.png")
        except:
            pyautogui.alert(text=f'Unable to find vista lists, make sure patient list is open and that the cancel and specialties buttons are visible then press Try Again', title='Unable to find Vista lists', button='Try Again')
    # Using screenshots to find the "Specialties" buttons    
    try:
        spec=pyautogui.locateOnScreen(fr"Data\{FONT_SIZE}\spec.png")
        pyautogui.click(spec.left,spec.top)
    except:
        print("Failed to find specialties button")
    # Program pauses for one second to ensure that the specialties list loaded. 
    # It then tries to find each team. If not found it tries to scroll down automatically
    # If patient is still ultimately not found it tells the user to manually locate it
    for team in teams.values():
        found = False
        # Tries to look for the team being counted based on the screenshot defined in the team dictionary
        # If it fails it attempts scrolling down to find team, if it still fails it will prompt the user to 
        # find it manually
        while not found: 
            try:
                box=pyautogui.locateOnScreen(team.image)
                pyautogui.click(box.left,box.top)
                found=True
            except:
                try:
                    box=pyautogui.locateOnScreen(fr"Data\{FONT_SIZE}\scroll.png")
                    pyautogui.click(box.left,box.top+20)
                    time.sleep(1)
                except:
                    pyautogui.alert(text=f'Program failed to scroll automatically to find {team.name}, please scroll manually until it is visible and click OK', title='Scroll failed', button='OK')
        team.patients=getPatientList(team.name=="Palliative")
    transferPatients()
    updateCounts()
    getLastExcel()
    comparePatients()
    writeExcel()
    getPersonal()
    writeWord()
        
    
def transferOnly():
    excelName= f'{date.today()}.xlsx'
    if os.path.isfile(excelName):
        wb = openpyxl.load_workbook(excelName)
        for team in teams.values():
            teamSheet=wb[team.name]
            team.patients.clear()
            r=2
            while True:
                name=teamSheet.cell(row=r,column=2).value
                room=teamSheet.cell(row=r,column=3).value
                mrn=teamSheet.cell(row=r,column=4).value
                if not mrn or str(mrn).isspace(): break
                team.patients.append(Patient(name,room,mrn))           
                r+=1
        wb.close()
        transferPatients()
        updateCounts()
        getLastExcel()
        comparePatients()
        writeExcel()
        writeWord()
    else:
        pyautogui.alert(text='No excel file found for today, please run Lists first', title='Error', button='OK')


    
    
window = tk.Tk()
greeting = ttk.Label(text="Please login to CPRS and open the select patient window, then press Begin when you are ready")
greeting.pack()
okButton = ttk.Button(text="Lists", command=records)
okButton.pack()
okButton = ttk.Button(text="Transfer Only", command=transferOnly)
okButton.pack()
window.text="Automatic Records Parser"
window.mainloop()